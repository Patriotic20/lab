"""
/// script
requires-python = ">=3.10"
dependencies = ["python-docx>=1.1.0"]
///

ndktu-student-face-platform loyihasining "Masofaviy ta'limni joriy etilishini
o'rganish (OTM)" davlat chek-listiga muvofiqligi to'g'risida o'zbek tilidagi
hisobotni (.docx) yaratuvchi skript.

Ishga tushirish:
    uv run --with python-docx python3 scripts/generate_compliance_report_uz.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent.parent / "Hisobot_LMS_OTM_muvofiqligi.docx"

STATUS_COLORS = {
    "Amalga oshirilgan": RGBColor(0x1F, 0x7A, 0x1F),
    "Qisman": RGBColor(0xC9, 0x7B, 0x00),
    "Mavjud emas": RGBColor(0xB0, 0x1B, 0x1B),
    "Tashkiliy": RGBColor(0x4A, 0x4A, 0x4A),
}


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:cs", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), "Times New Roman")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_paragraph(doc: Document, text: str, bold: bool = False,
                  italic: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Hisobot mazmuni
# ---------------------------------------------------------------------------

CHECKLIST: list[tuple[str, str, str, str, str]] = [
    (
        "1. Masofaviy ta'lim markazi",
        "1",
        "Masofaviy ta'limni boshqarish markazi (alohida tarkibiy bo'linma) tashkil etilganligi",
        "Tashkiliy",
        "Dasturiy qismga tegishli emas. Rektor buyrug'i bilan hal qilinadi.",
    ),
    (
        "2. Masofaviy ta'lim tamoyillari va talablari",
        "1",
        "Ishtirokchilar interaktivligi (muloqot, savol-javob, muhokama)",
        "Qisman",
        "Test va so'rovnomalar orqali amalga oshirilgan (quiz, question, user_answers "
        "modullari). Onlayn rejimda chat, forum va jonli muhokamalar mavjud emas.",
    ),
    (
        "2. Masofaviy ta'lim tamoyillari va talablari",
        "2",
        "Talabalarni aniqlash va autentifikatsiya qilish (IIV/GSP + HEMIS)",
        "Qisman",
        "JWT-autentifikatsiya (backend/app/modules/user), HEMIS bilan integratsiya "
        "(backend/app/modules/hemis/service.py), yuzni biometrik tasdiqlash "
        "(face-detection/app/services/face_detector.py — MediaPipe + face_recognition). "
        "IIV/GSP bilan to'g'ridan-to'g'ri integratsiya yo'q.",
    ),
    (
        "2. Masofaviy ta'lim tamoyillari va talablari",
        "3",
        "Virtual sinf (real vaqtdagi onlayn darslar, yozib olish, LMS bilan integratsiya)",
        "Mavjud emas",
        "To'liq virtual sinf amalga oshirilmagan. Video kanallar faqat imtihonlarni "
        "nazorat qilish uchun ishlatiladi (face-detection xizmati).",
    ),
    (
        "2. Masofaviy ta'lim tamoyillari va talablari",
        "4",
        "Umumiy, guruh va individual ta'lim shakllarining uyg'unligi",
        "Qisman",
        "Ma'ruzalar (lesson moduli), topshiriq va testlar (quiz, question), "
        "resurslar (resource) qo'llab-quvvatlanadi. Seminarlar va individual "
        "konsultatsiyalarni tizimli o'tkazish vositasi yo'q.",
    ),
    (
        "3. Moddiy-texnik baza",
        "1",
        "LMS platformasining mavjudligi (Learning Management System)",
        "Amalga oshirilgan",
        "Platforma joriy etilgan: FastAPI backend + React 19 SPA + PostgreSQL 17 + Redis. "
        "To'liq konteynerlashtirilgan (docker-compose.yml).",
    ),
    (
        "3. Moddiy-texnik baza",
        "2",
        "Videodarslarni ishlab chiqish uchun studiya",
        "Mavjud emas",
        "Video-ma'ruzalarni yuklash, transkodlash va translatsiya qilish quyi tizimlari "
        "yo'q. Fayl yuklash imkoniyati faqat resource modulida mavjud.",
    ),
    (
        "3. Moddiy-texnik baza",
        "3",
        "O'quv yili uchun barcha fanlar bo'yicha o'quv kontentini mavjudligi",
        "Qisman",
        "Resource moduli (backend/app/modules/resource) fayllar va havolalarni saqlaydi. "
        "O'quv kalendari va fanlar bo'yicha to'plamlar bilan qat'iy bog'lanish yo'q.",
    ),
    (
        "3. Moddiy-texnik baza",
        "4",
        "Barcha fanlar bo'yicha elektron O'MM va raqamli kutubxona",
        "Qisman",
        "Generic resource moduli amalga oshirilgan. Tarkibiy O'MM (uslubiy ko'rsatmalar, "
        "ishchi dasturlar, sillabuslar) va ilmiy adabiyotlarning raqamli kutubxonasi yo'q.",
    ),
    (
        "3. Moddiy-texnik baza",
        "5",
        "Malakali muhandis-texnik xodimlar",
        "Tashkiliy",
        "Dasturiy qismga tegishli emas.",
    ),
    (
        "3. Moddiy-texnik baza",
        "6",
        "Server infratuzilmasi (O'zbekiston hududida)",
        "Tashkiliy",
        "Konteynerlar tayyor (docker-compose.yml: PostgreSQL, Redis, FastAPI, React, "
        "face-detection, Grafana). Uskunalarni joylashtirish — tashkiliy masala.",
    ),
    (
        "3. Moddiy-texnik baza",
        "7",
        "OTM rasmiy veb-sayti (nizom, o'quv rejalari, PPS, akademik kalendar)",
        "Mavjud emas",
        "Frontend administrator, o'qituvchi va talabalar uchun ichki dashboard hisoblanadi. "
        "OTM ning ommaviy veb-sahifasi yo'q.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "1",
        "LMS ↔ HEMIS integratsiyasi",
        "Amalga oshirilgan",
        "backend/app/modules/hemis/service.py da to'liq integratsiya: avtorizatsiya, "
        "fakultetlar, guruhlar va talabalarni HEMIS API orqali sinxronlash.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "2",
        "Avtoproktoring (kamera, ekran, AI monitoring)",
        "Qisman",
        "face-detection mikroservisi (face-detection/app/services/video_service.py): "
        "MediaPipe orqali yuzni aniqlash va bir nechta yuzlarni topish. "
        "Eye-tracking, head-pose estimation, ekran monitoringi va vkladkalarni "
        "almashtirish detektsiyasi yo'q.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "3",
        "Axborot resurslari komponenti (video, matn, fayl, test)",
        "Qisman",
        "Resource moduli (backend/app/modules/resource): matnlar, fayllar, havolalar. "
        "Video-translatsiya va DRM himoyasi yo'q.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "4",
        "Boshqarish komponenti (admin panel)",
        "Amalga oshirilgan",
        "SQLAdmin (backend/app/main.py:70-72) + RBAC: role, permission, user modullari. "
        "Foydalanuvchilar va rollarni to'liq boshqarish.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "5",
        "Davomat va o'zlashtirishni hisobga olish komponenti",
        "Qisman",
        "LessonResult modelida attendance maydoni mavjud (backend/app/modules/lesson/model.py), "
        "umumiy statistika (backend/app/modules/statistics). Har bir dars bo'yicha "
        "to'liq davomat jurnali mavjud emas.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "6",
        "Kommunikatsiya komponenti (chat, forum, xabarlar)",
        "Mavjud emas",
        "O'qituvchilar va talabalar o'rtasidagi ichki muloqot vositalari yo'q.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "7",
        "Kontingentni hisobga olish komponenti (talabalar, guruhlar, harakat)",
        "Amalga oshirilgan",
        "Iyerarxiya: faculty → kafedra → group → student "
        "(backend/app/modules/{faculty,kafedra,group,student}). "
        "HEMIS bilan bog'lanish ma'lumotlarning dolzarbligini ta'minlaydi.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "8",
        "Kurslarni boshqarish komponenti",
        "Amalga oshirilgan",
        "subject, subject_teacher, lesson modullari: kurslarni yaratish, tahrirlash, "
        "tuzilmalashtirish va talabalarni guruhlar orqali biriktirish.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "9",
        "O'qitishni boshqarish komponenti (rejalashtirish, topshiriqlar, monitoring)",
        "Qisman",
        "quiz, quiz_process modullari: testlar, urinishlar, davomiyligi. To'liq "
        "kalendar rejalashtirish va dars jadvali yo'q.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "10",
        "Statistika va analitika komponenti",
        "Amalga oshirilgan",
        "statistics moduli (backend/app/modules/statistics): umumiy, test, foydalanuvchi, "
        "fakultet, guruh va o'qituvchi statistikasi. Grafana dashboardlari.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "11",
        "Talabalar bilimini nazorat qilish komponenti (test, imtihon, topshiriq)",
        "Amalga oshirilgan",
        "quiz, question, user_answers, result modullari. Bir nechta savol turlarini "
        "qo'llab-quvvatlash, avtomatik baholash, ko'p marotaba urinishlar.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "12",
        "Yagona reestrda (reestr.uz) ro'yxatdan o'tkazilganligi",
        "Tashkiliy",
        "Dasturiy qismga tegishli emas. Ariza topshirish va xulosa olish talab etiladi.",
    ),
    (
        "4. LMS platformasiga qo'yilgan talablar",
        "13",
        "«Kiberxavfsizlik markazi» (DUK) xulosasi",
        "Tashkiliy",
        "Dasturiy qismga tegishli emas. «Kiberxavfsizlik markazi» davlat unitar korxonasining "
        "mustaqil ekspertizasini talab qiladi.",
    ),
]


IMPLEMENTED_DETAILS = [
    ("Autentifikatsiya va avtorizatsiya (JWT + RBAC)",
     "JWT tokenlar va refresh-tokenlar, localStorage da saqlanadi; tokenlarni shaffof "
     "yangilash uchun axios-interseptor (frontend/src/services/api.ts). "
     "Rollar: admin, teacher, student, psixologik, tutor — role, permission, user modullari."),
    ("HEMIS bilan integratsiya",
     "hemis moduli (backend/app/modules/hemis): avtorizatsiya, fakultetlar, guruhlar va "
     "talabalarni HEMIS davlat tizimi bilan sinxronlash."),
    ("Yuzni biometrik tasdiqlash",
     "face-detection mikroservisi (face-detection/app/services/face_detector.py): "
     "MediaPipe + face_recognition; imtihon boshlanishida biometriyani solishtirish."),
    ("Testlar va bilimni nazorat qilish",
     "quiz, question, user_answers, result modullari: davomiylik va urinishlar soni bilan "
     "testlar yaratish, bir nechta savol turlarini qo'llab-quvvatlash, avtomatik baholash."),
    ("Kontingentni hisobga olish",
     "Iyerarxik tuzilma: faculty → kafedra → group → student. Talabalarni guruhlarga, "
     "o'qituvchilarni kafedralar va fanlarga biriktirish."),
    ("Kurslar va darslarni boshqarish",
     "subject, subject_teacher, lesson modullari: dastur, o'qituvchi, dars sanasi."),
    ("Psixologiya moduli",
     "psychology mustaqil quyi tizimi: metodlar, savollar (6 ta tur), avtomatik ball "
     "hisoblash (psychology/scoring.py)."),
    ("Statistika va monitoring",
     "statistics moduli + Grafana dashboardlari (nusmt_grafana). "
     "Foydalanuvchilar, guruhlar, fakultetlar va o'qituvchilar bo'yicha hisobotlar."),
    ("Fayl resurslari",
     "resource moduli: fayllarni yuklash va havolalarni saqlash, statik fayllarni "
     "/uploads orqali tarqatish."),
    ("Admin-panel",
     "SQLAdmin (backend/app/main.py): RBAC tekshiruvi bilan barcha mavjudotlar bo'yicha "
     "to'liq CRUD interfeys."),
    ("Konteynerlashtirish va joylashtirish",
     "docker-compose.yml: 6 ta xizmat (backend, frontend, face-detection, postgres, "
     "redis, grafana). Health-check, make deploy orqali zero-downtime joylashtirish."),
]


MISSING_LIST = [
    "Real vaqtdagi onlayn ma'ruzalarni o'tkazish va yozib olish imkoniyati bilan "
    "virtual sinf (BBB/Jitsi/WebRTC).",
    "O'qituvchilar va talabalar o'rtasidagi chat, forum va shaxsiy xabarlar.",
    "OTM ning ommaviy veb-sayti: nizom, o'quv rejalari, PPS tarkibi, akademik kalendar, "
    "yangiliklar, kontaktlar.",
    "Videostudiya: video-ma'ruzalarni yuklash, transkodlash (ffmpeg) va adaptiv "
    "translatsiya qilish (HLS/DASH).",
    "Barcha fanlar bo'yicha elektron o'quv-uslubiy majmualar (O'MM) va raqamli kutubxona.",
    "Dars jadvali va kontentni o'quv kalendari bilan bog'lash.",
    "Har bir dars bo'yicha to'liq elektron davomat jurnali (faqat test natijalari emas).",
    "Kengaytirilgan proktoring: eye-tracking, head-pose, vkladkalarni almashtirish "
    "detektsiyasi, audio monitoring, AI-anomaliyalar.",
    "Pasport ma'lumotlarini tekshirish uchun IIV (IIB) / GSP tizimi bilan to'g'ridan-to'g'ri integratsiya.",
    "LMS ni «Raqamli hukumatning yagona reestri» (reestr.uz) da ro'yxatdan o'tkazish va "
    "ijobiy xulosa olish.",
    "«Kiberxavfsizlik markazi» (DUK) ekspertizasidan o'tish.",
]


RECOMMENDATIONS = [
    ("Virtual sinf",
     "Jitsi Meet yoki BigBlueButton ni WebRTC orqali integratsiya qilish. Ma'ruzalarni "
     "S3-mos saqlash (MinIO) ga yozib olish. classroom modulini yaratish: jadval, "
     "uchrashuv havolalari, yozuvlarni LMS ga avtomatik yuklash. Baholash: 4–6 odam-haftasi."),
    ("Proktoringni kengaytirish",
     "face-detection ga qo'shish: MediaPipe Iris (eye-tracking), head-pose uchun PnP-algoritm, "
     "document.visibilitychange JS hodisasi (vkladka almashishni aniqlash), atrof-muhit "
     "tovushini tahlil qilish. Xulq-atvor anomaliyalari uchun AI-klassifikator. "
     "Baholash: 3–4 odam-haftasi."),
    ("Chat, forum, xabarlar",
     "FastAPI + Redis pub/sub asosida WebSocket-chat. Fanlar bo'yicha mavzuli forumlar. "
     "Rollar o'rtasidagi shaxsiy xabarlar. WebPush orqali bildirishnomalar. "
     "Baholash: 3 odam-haftasi."),
    ("OTM ning ommaviy veb-sayti",
     "Alohida Next.js yoki React-SPA ilovasi. Bo'limlar: «OTM haqida», «Nizom», «PPS», "
     "«O'quv rejalari», «Akademik kalendar», «Yangiliklar», «Qabul», «Kontaktlar». "
     "Mavjud backend ga ommaviy read-only endpointlar orqali ulanish. "
     "Baholash: 4–5 odam-haftasi."),
    ("Elektron kutubxona va O'MM",
     "O'MM ning tarkibiy katalogi: ishchi dasturlar, sillabuslar, uslubiy ko'rsatmalar, "
     "laboratoriyalar. Brauzerda DRM-ko'rishi (PDF.js + watermarking). Fanlar va "
     "semestrlar bo'yicha kategoriyalash. Baholash: 3–4 odam-haftasi."),
    ("Elektron davomat jurnali",
     "Alohida attendance moduli: har bir dars bo'yicha qatnashishni qayd etish. "
     "Rejalashtirilgan darslar bilan bog'lanish. Virtual sinfga kirish vaqtida "
     "avtomatik belgi qo'yish. Dekanat uchun hisobotlar. Baholash: 2 odam-haftasi."),
    ("Videostudiya va translatsiya",
     "Pipeline: yuklash → ffmpeg-transkodlash (1080p/720p/480p) → HLS-manifest → "
     "CDN orqali yetkazib berish. Video metama'lumotlarni tahrirlash interfeysi. "
     "Baholash: 3–4 odam-haftasi."),
    ("Dars jadvali",
     "schedule moduli: semestrlar bo'yicha darslar timetable i. Dashboardda kalendar "
     "vidjeti. iCal-eksport. Baholash: 2 odam-haftasi."),
    ("IIV/GSP bilan integratsiya",
     "Ro'yxatdan o'tishda pasport ma'lumotlarini tekshirish uchun IIV API ga ulanish. "
     "Rasmiy ruxsat olishni talab qiladi. Baholash: 2 odam-haftasi + tashkiliy muddatlar."),
    ("reestr.uz da ro'yxatdan o'tish va DUK xulosasi",
     "Hujjatlar to'plamini tayyorlash, kiberxavfsizlik ekspertizasidan o'tish, kamchiliklarni "
     "bartaraf etish, ijobiy xulosa olish. Muddat: 2–4 oy (tashkiliy)."),
]


ROADMAP = [
    ("P0 — OTM ga muvofiqlik uchun kritik",
     [
         "Virtual sinf (4–6 hafta).",
         "Proktoringni kengaytirish (3–4 hafta).",
         "Chat va forum (3 hafta).",
         "reestr.uz da ro'yxatdan o'tish va DUK xulosasi (tashkiliy, 2–4 oy).",
     ]),
    ("P1 — o'rta muddatda zarur",
     [
         "OTM ning ommaviy veb-sayti (4–5 hafta).",
         "Elektron kutubxona va O'MM (3–4 hafta).",
         "Elektron davomat jurnali (2 hafta).",
         "Videostudiya va translatsiya (3–4 hafta).",
     ]),
    ("P2 — istalgan yaxshilanishlar",
     [
         "Dars jadvali (2 hafta).",
         "IIV/GSP bilan integratsiya (2 hafta).",
     ]),
]


# ---------------------------------------------------------------------------
# Hujjatni yig'ish
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Sarlavha
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HISOBOT\nmasofaviy ta'lim platformasining\n"
                        "«Masofaviy ta'limni joriy etilishini o'rganish (OTM)»\n"
                        "davlat chek-listiga muvofiqligi to'g'risida")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    add_paragraph(doc, "")
    add_paragraph(doc, "Platforma: ndktu-student-face-platform",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Tayyorlangan sana: 2026-yil 10-may",
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "")

    # 1. Kirish
    add_heading(doc, "1. Kirish", level=1)
    add_paragraph(
        doc,
        "Ushbu hisobot ndktu-student-face-platform dasturiy platformasining "
        "«Masofaviy ta'limni joriy etilishini o'rganish (OTM)» chek-listida "
        "bayon etilgan talablarga muvofiqligini baholash maqsadida tayyorlangan. "
        "Hujjatda 25 ta ko'rsatkich bo'yicha muvofiqlikning umumiy jadvali, "
        "amalga oshirilgan funksionallikning tavsifi, mavjud bo'lmagan "
        "komponentlar ro'yxati, shuningdek, takomillashtirish bo'yicha tavsiyalar "
        "va prioritetlangan yo'l xaritasi keltirilgan."
    )
    add_paragraph(
        doc,
        "Platforma arxitekturasi quyidagi texnologiyalar to'plamiga asoslanadi:"
    )
    add_bullets(doc, [
        "Backend: FastAPI (Python 3.12, async), SQLAlchemy 2, PostgreSQL 17, Alembic.",
        "Frontend: React 19 + Vite + TypeScript + Tailwind CSS 4, React Query, Zod.",
        "Yuzni aniqlash mikroservisi: MediaPipe + face_recognition asosidagi alohida FastAPI xizmati.",
        "Kesh va navbatlar: Redis 7.",
        "Monitoring: Prometheus + Grafana.",
        "Joylashtirish: Docker Compose, nginx-proksi, zero-downtime deploy.",
        "Avtorizatsiya: JWT (access + refresh), role / permission modullari orqali RBAC.",
        "Admin-panel: SQLAdmin.",
    ])

    # 2. Muvofiqlik jadvali
    add_heading(doc, "2. Muvofiqlikning umumiy jadvali", level=1)
    add_paragraph(
        doc,
        "Status rang belgilari: yashil — amalga oshirilgan, to'q sariq — qisman "
        "amalga oshirilgan, qizil — mavjud emas, kulrang — tashkiliy tadbirlarga taalluqli."
    )

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.5), Cm(0.9), Cm(5.0), Cm(2.6), Cm(5.0)]
    hdr = table.rows[0].cells
    headers = ["Bo'lim", "№", "Ko'rsatkich", "Status", "Izoh"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        set_cell_shading(hdr[i], "DDDDDD")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr[i].width = widths[i]

    for section_name, num, indicator, status, comment in CHECKLIST:
        row = table.add_row().cells
        for i, value in enumerate([section_name, num, indicator, status, comment]):
            row[i].text = ""
            p = row[i].paragraphs[0]
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            row[i].width = widths[i]
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i == 3:
                color = STATUS_COLORS.get(status)
                if color:
                    run.bold = True
                    run.font.color.rgb = color
                fill_map = {
                    "Amalga oshirilgan": "DFF5DF",
                    "Qisman": "FCEAB6",
                    "Mavjud emas": "F8D7D7",
                    "Tashkiliy": "ECECEC",
                }
                if status in fill_map:
                    set_cell_shading(row[i], fill_map[status])

    # 3. Amalga oshirilgan funksiyalar
    add_heading(doc, "3. Amalga oshirilgan funksiyalar", level=1)
    add_paragraph(
        doc,
        "Quyida platformaning joriy versiyasida allaqachon amalga oshirilgan va "
        "foydalanish uchun mavjud bo'lgan funksional komponentlarning batafsil "
        "ro'yxati keltirilgan."
    )
    for idx, (name, desc) in enumerate(IMPLEMENTED_DETAILS, start=1):
        add_heading(doc, f"3.{idx}. {name}", level=2)
        add_paragraph(doc, desc)

    # 4. Mavjud bo'lmagan funksiyalar
    add_heading(doc, "4. Mavjud bo'lmagan funksiyalar", level=1)
    add_paragraph(
        doc,
        "Kod bazasini auditidan so'ng OTM talablariga nisbatan quyidagi funksional "
        "kamchiliklar aniqlandi:"
    )
    add_bullets(doc, MISSING_LIST)

    # 5. Tavsiyalar
    add_heading(doc, "5. Takomillashtirish bo'yicha tavsiyalar", level=1)
    add_paragraph(
        doc,
        "OTM talablariga to'liq muvofiq bo'lish uchun quyidagi takomillashtirishlarni "
        "amalga oshirish taklif etiladi. Har bir tashabbus uchun tavsiya etilgan "
        "texnologiyalar va taxminiy mehnat sarfi ko'rsatilgan."
    )
    for idx, (name, desc) in enumerate(RECOMMENDATIONS, start=1):
        add_heading(doc, f"5.{idx}. {name}", level=2)
        add_paragraph(doc, desc)

    # 6. Yo'l xaritasi
    add_heading(doc, "6. Yo'l xaritasi", level=1)
    add_paragraph(
        doc,
        "Takomillashtirishlar prioritetlar bo'yicha guruhlangan. Ishlarning umumiy "
        "hajmi — taxminan 26–34 odam-haftasi va tashkiliy tadbirlar."
    )
    for prio_name, items in ROADMAP:
        add_heading(doc, prio_name, level=2)
        add_bullets(doc, items)

    # 7. Xulosa
    add_heading(doc, "7. Xulosa", level=1)
    add_paragraph(
        doc,
        "ndktu-student-face-platform platformasi OTM chek-listi talablarining "
        "muhim qismini qoplaydi: LMS o'zagi, HEMIS bilan integratsiya, biometrik "
        "identifikatsiya, RBAC, bilimni nazorat qilish, statistika va kontingentni "
        "hisobga olish amalga oshirilgan. To'liq muvofiqlik uchun virtual sinf, "
        "kengaytirilgan proktoring, kommunikatsiya vositalari, OTM ning ommaviy "
        "veb-sayti va elektron kutubxonani ishlab chiqish, shuningdek, tashkiliy "
        "tadbirlardan o'tish (reestr.uz da ro'yxatdan o'tish va DUK ekspertizasi) talab etiladi."
    )
    add_paragraph(
        doc,
        "P0 prioritetidagi ishlarni boshlash va parallel ravishda Yagona reestrda "
        "ro'yxatdan o'tish va kiberxavfsizlik ekspertizasi uchun hujjatlarni "
        "rasmiylashtirishni boshlash tavsiya etiladi."
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Hisobot saqlandi: {OUTPUT}")


if __name__ == "__main__":
    main()
