"""
/// script
requires-python = ">=3.10"
dependencies = ["python-docx>=1.1.0"]
///

Generate a Russian-language compliance report (.docx) mapping the
ndktu-student-face-platform features against the Uzbek state checklist
"Masofaviy ta'limni joriy etilishini o'rganish (OTM)".

Run:
    uv run --with python-docx scripts/generate_compliance_report.py
or:
    python3 scripts/generate_compliance_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent.parent / "Otchet_o_sootvetstvii_LMS_OTM.docx"

STATUS_COLORS = {
    "Реализовано": RGBColor(0x1F, 0x7A, 0x1F),
    "Частично": RGBColor(0xC9, 0x7B, 0x00),
    "Отсутствует": RGBColor(0xB0, 0x1B, 0x1B),
    "Организационный": RGBColor(0x4A, 0x4A, 0x4A),
}


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:cs", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), "Times New Roman")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_paragraph(doc: Document, text: str, bold: bool = False,
                  italic: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Контент отчёта
# ---------------------------------------------------------------------------

CHECKLIST: list[tuple[str, str, str, str, str]] = [
    # (Раздел, №, Показатель, Статус, Комментарий)
    (
        "1. Центр дистанционного обучения",
        "1",
        "Создание Центра дистанционного обучения (отдельная структурная единица)",
        "Организационный",
        "Не относится к программной части. Решается приказом ректора.",
    ),
    (
        "2. Принципы и требования к организации ДО",
        "1",
        "Интерактивность участников (диалог, обсуждения, Q&A)",
        "Частично",
        "Реализовано через викторины и тесты (модули quiz, question, user_answers). "
        "Отсутствуют чат, форум, реальное обсуждение в режиме онлайн.",
    ),
    (
        "2. Принципы и требования к организации ДО",
        "2",
        "Идентификация и аутентификация студентов (IIV/GSP + HEMIS)",
        "Частично",
        "JWT-аутентификация (backend/app/modules/user), интеграция с HEMIS "
        "(backend/app/modules/hemis/service.py), биометрическая верификация лица "
        "(face-detection/app/services/face_detector.py — MediaPipe + face_recognition). "
        "Отсутствует прямая интеграция с IIV/GSP для проверки паспортных данных.",
    ),
    (
        "2. Принципы и требования к организации ДО",
        "3",
        "Виртуальный класс (онлайн-уроки в реальном времени, запись, интеграция с LMS)",
        "Отсутствует",
        "Полноценный виртуальный класс не реализован. Видеоканалы используются только "
        "для прокторинга экзаменов (face-detection сервис).",
    ),
    (
        "2. Принципы и требования к организации ДО",
        "4",
        "Согласованность общих, групповых и индивидуальных форм обучения",
        "Частично",
        "Поддержаны лекции (модуль lesson), задания и тесты (quiz, question), "
        "ресурсы (resource). Отсутствует структурированное проведение семинаров "
        "и индивидуальных консультаций.",
    ),
    (
        "3. Материально-техническая база",
        "1",
        "Наличие LMS-платформы (Learning Management System)",
        "Реализовано",
        "Платформа развёрнута: FastAPI backend + React 19 SPA + PostgreSQL 17 + Redis. "
        "Полная контейнеризация (docker-compose.yml).",
    ),
    (
        "3. Материально-техническая база",
        "2",
        "Студия для разработки видеоуроков",
        "Отсутствует",
        "Подсистемы загрузки, транскодинга и стриминга видео-лекций нет. "
        "Поле для загрузки файлов есть только в модуле resource.",
    ),
    (
        "3. Материально-техническая база",
        "3",
        "Учебный контент на учебный год по всем дисциплинам",
        "Частично",
        "Модуль resource (backend/app/modules/resource) хранит файлы и ссылки. "
        "Отсутствует строгая привязка к учебному календарю и комплектам по дисциплинам.",
    ),
    (
        "3. Материально-техническая база",
        "4",
        "Электронные УМК и цифровая библиотека по всем дисциплинам",
        "Частично",
        "Реализован generic-модуль resource. Структурированные УМК "
        "(методические указания, рабочие программы, силлабусы) и цифровая библиотека "
        "научной литературы отсутствуют.",
    ),
    (
        "3. Материально-техническая база",
        "5",
        "Квалифицированный инженерно-технический персонал",
        "Организационный",
        "Не относится к программной части.",
    ),
    (
        "3. Материально-техническая база",
        "6",
        "Серверная инфраструктура (на территории Узбекистана)",
        "Организационный",
        "Контейнеры готовы (docker-compose.yml: PostgreSQL, Redis, FastAPI, React, "
        "face-detection, Grafana). Размещение оборудования — организационный вопрос.",
    ),
    (
        "3. Материально-техническая база",
        "7",
        "Официальный веб-сайт ВУЗа (устав, учебные планы, ППС, академический календарь)",
        "Отсутствует",
        "Frontend представляет собой внутренний дашборд для администраторов, "
        "преподавателей и студентов. Публичной части сайта ВУЗа нет.",
    ),
    (
        "4. Требования к LMS-платформе",
        "1",
        "Интеграция LMS ↔ HEMIS",
        "Реализовано",
        "Полная интеграция в backend/app/modules/hemis/service.py: авторизация, "
        "синхронизация факультетов, групп и студентов через HEMIS API.",
    ),
    (
        "4. Требования к LMS-платформе",
        "2",
        "Автопрокторинг (камера, экран, AI-мониторинг)",
        "Частично",
        "Микросервис face-detection (face-detection/app/services/video_service.py): "
        "распознавание лица, обнаружение нескольких лиц через MediaPipe. "
        "Отсутствуют eye-tracking, head-pose estimation, мониторинг экрана и "
        "детекция переключения вкладок.",
    ),
    (
        "4. Требования к LMS-платформе",
        "3",
        "Компонент информационных ресурсов (видео, текст, файл, тест)",
        "Частично",
        "Модуль resource (backend/app/modules/resource): тексты, файлы, ссылки. "
        "Видео-стриминга и DRM-защиты нет.",
    ),
    (
        "4. Требования к LMS-платформе",
        "4",
        "Компонент управления (admin panel)",
        "Реализовано",
        "SQLAdmin (backend/app/main.py:70-72) + RBAC: модули role, permission, user. "
        "Полное управление пользователями и ролями.",
    ),
    (
        "4. Требования к LMS-платформе",
        "5",
        "Компонент учёта посещаемости и успеваемости",
        "Частично",
        "Поле attendance в модели LessonResult (backend/app/modules/lesson/model.py), "
        "общая статистика (backend/app/modules/statistics). Отсутствует полноценный "
        "журнал посещаемости по каждому занятию.",
    ),
    (
        "4. Требования к LMS-платформе",
        "6",
        "Компонент коммуникации (чат, форум, сообщения)",
        "Отсутствует",
        "Внутренних коммуникационных средств между преподавателями и студентами нет.",
    ),
    (
        "4. Требования к LMS-платформе",
        "7",
        "Компонент учёта контингента (студенты, группы, движение)",
        "Реализовано",
        "Иерархия factuality → kafedra → group → student "
        "(backend/app/modules/{faculty,kafedra,group,student}). "
        "Связь с HEMIS обеспечивает актуальность данных.",
    ),
    (
        "4. Требования к LMS-платформе",
        "8",
        "Компонент управления курсами",
        "Реализовано",
        "Модули subject, subject_teacher, lesson: создание, редактирование, "
        "структурирование курсов и привязка студентов через группы.",
    ),
    (
        "4. Требования к LMS-платформе",
        "9",
        "Компонент управления обучением (планирование, задания, мониторинг)",
        "Частично",
        "Модули quiz, quiz_process: викторины, попытки, длительность. Отсутствует "
        "полное календарное планирование и расписание занятий.",
    ),
    (
        "4. Требования к LMS-платформе",
        "10",
        "Компонент статистики и аналитики",
        "Реализовано",
        "Модуль statistics (backend/app/modules/statistics): общая, поквиктовая, "
        "пользовательская, факультетская, групповая и преподавательская статистика. "
        "Дашборды в Grafana.",
    ),
    (
        "4. Требования к LMS-платформе",
        "11",
        "Компонент контроля знаний (тесты, экзамены, задания)",
        "Реализовано",
        "Модули quiz, question, user_answers, result. Поддержка нескольких типов "
        "вопросов, автоматическая оценка, многократные попытки.",
    ),
    (
        "4. Требования к LMS-платформе",
        "12",
        "Регистрация в Едином реестре (reestr.uz)",
        "Организационный",
        "Не относится к программной части. Требует подачи заявки и заключения.",
    ),
    (
        "4. Требования к LMS-платформе",
        "13",
        "Заключение «Кибербезопасности» (DUK)",
        "Организационный",
        "Не относится к программной части. Требует независимой экспертизы "
        "Государственного унитарного предприятия «Центр кибербезопасности».",
    ),
]


IMPLEMENTED_DETAILS = [
    ("Аутентификация и авторизация (JWT + RBAC)",
     "JWT-токены и refresh-токены, хранение в localStorage; интерцептор axios "
     "для прозрачного обновления токенов (frontend/src/services/api.ts). "
     "Роли: admin, teacher, student, psixologik, tutor — модели role, permission, user."),
    ("Интеграция с HEMIS",
     "Модуль hemis (backend/app/modules/hemis): авторизация, синхронизация "
     "факультетов, групп и студентов с государственной системой HEMIS."),
    ("Биометрическая верификация лица",
     "Микросервис face-detection (face-detection/app/services/face_detector.py): "
     "MediaPipe + face_recognition; сравнение биометрии при запуске экзамена."),
    ("Викторины и контроль знаний",
     "Модули quiz, question, user_answers, result: создание тестов с длительностью "
     "и количеством попыток, поддержка нескольких типов вопросов, автоматическая оценка."),
    ("Учёт контингента",
     "Иерархическая структура: faculty → kafedra → group → student. Привязка "
     "студентов к группам, преподавателей к кафедрам и предметам."),
    ("Управление курсами и занятиями",
     "Модули subject, subject_teacher, lesson: программа, преподаватель, дата занятия."),
    ("Психологический модуль",
     "Самостоятельная подсистема psychology: методы, вопросы (6 типов), "
     "автоматический подсчёт баллов (psychology/scoring.py)."),
    ("Статистика и мониторинг",
     "Модуль statistics + Grafana-дашборды (nusmt_grafana). "
     "Отчёты по пользователям, группам, факультетам, преподавателям."),
    ("Файловые ресурсы",
     "Модуль resource: загрузка файлов и сохранение ссылок, статическая раздача через /uploads."),
    ("Админ-панель",
     "SQLAdmin (backend/app/main.py): полный CRUD-интерфейс по всем сущностям с "
     "RBAC-проверкой."),
    ("Контейнеризация и развёртывание",
     "docker-compose.yml: 6 сервисов (backend, frontend, face-detection, postgres, "
     "redis, grafana). Health-checks, zero-downtime deploy через make deploy."),
]


MISSING_LIST = [
    "Виртуальный класс с возможностью проведения онлайн-лекций в реальном времени и записи (BBB/Jitsi/WebRTC).",
    "Чат, форум и личные сообщения между преподавателями и студентами.",
    "Публичный сайт ВУЗа: устав, учебные планы, состав ППС, академический календарь, новости, контакты.",
    "Видеостудия: загрузка, транскодинг (ffmpeg) и адаптивный стриминг (HLS/DASH) видеолекций.",
    "Электронные учебно-методические комплексы (УМК) и цифровая библиотека по всем дисциплинам.",
    "Календарное расписание занятий и привязка контента к учебному календарю.",
    "Полноценный электронный журнал посещаемости (по каждому занятию, не только по результатам теста).",
    "Расширенный прокторинг: eye-tracking, head-pose, детекция переключения вкладок, "
    "звуковой мониторинг, AI-аномалии.",
    "Прямая интеграция с системой ИИВ (МВД) / GSP для верификации паспортных данных абитуриента.",
    "Регистрация LMS в едином реестре «Раздан хукумат» (reestr.uz) и получение положительного заключения.",
    "Прохождение экспертизы Государственного унитарного предприятия «Центр кибербезопасности» (DUK).",
]


RECOMMENDATIONS = [
    ("Виртуальный класс",
     "Интеграция Jitsi Meet или BigBlueButton через WebRTC. Запись лекций в S3-совместимое "
     "хранилище (MinIO). Создать модуль classroom: расписание, ссылки на встречи, "
     "автозагрузка записей в LMS. Оценка: 4–6 человеко-недель."),
    ("Расширение прокторинга",
     "Добавить в face-detection: MediaPipe Iris (eye-tracking), PnP-алгоритм для head-pose, "
     "JS-обработчик document.visibilitychange (детекция переключения вкладок), "
     "анализ окружающего звука. AI-классификатор аномалий поведения. Оценка: 3–4 человеко-недели."),
    ("Чат, форум, сообщения",
     "WebSocket-чат на FastAPI + Redis pub/sub. Тематические форумы по дисциплинам. "
     "Прямые сообщения между ролями. Уведомления через WebPush. Оценка: 3 человеко-недели."),
    ("Публичный сайт ВУЗа",
     "Отдельное Next.js или React-SPA приложение. Разделы: «О ВУЗе», «Устав», «ППС», "
     "«Учебные планы», «Академический календарь», «Новости», «Поступление», «Контакты». "
     "Подключение к существующему backend через публичные read-only endpoints. "
     "Оценка: 4–5 человеко-недель."),
    ("Электронная библиотека и УМК",
     "Структурированный каталог УМК: рабочие программы, силлабусы, методические указания, "
     "лабораторные. DRM-просмотр PDF в браузере (PDF.js + watermarking). Категоризация "
     "по дисциплинам и семестрам. Оценка: 3–4 человеко-недели."),
    ("Электронный журнал посещаемости",
     "Отдельный модуль attendance: запись посещения по каждому занятию. Связь со "
     "scheduled lessons. Автоотметка по факту входа в виртуальный класс. Отчёты для "
     "деканата. Оценка: 2 человеко-недели."),
    ("Видеостудия и стриминг",
     "Pipeline: загрузка → ffmpeg-транскодинг (1080p/720p/480p) → HLS-манифест → "
     "доставка через CDN. Интерфейс редактирования метаданных видео. "
     "Оценка: 3–4 человеко-недели."),
    ("Календарное расписание",
     "Модуль schedule: таймтаблица занятий по семестрам. Календарный виджет на дашборде. "
     "iCal-экспорт. Оценка: 2 человеко-недели."),
    ("Интеграция с ИИВ/GSP",
     "Подключение к API ИИВ для проверки паспортных данных при регистрации. "
     "Требует получения официального доступа. Оценка: 2 человеко-недели + организационные сроки."),
    ("Регистрация в reestr.uz и заключение DUK",
     "Подготовка пакета документов, прохождение экспертизы кибербезопасности, "
     "устранение замечаний, получение положительного заключения. "
     "Срок: 2–4 месяца (организационный)."),
]


ROADMAP = [
    ("P0 — критичные для соответствия OTM",
     [
         "Виртуальный класс (4–6 нед.).",
         "Расширение прокторинга (3–4 нед.).",
         "Чат и форум (3 нед.).",
         "Регистрация в reestr.uz и заключение DUK (организационно, 2–4 мес.).",
     ]),
    ("P1 — необходимые в среднесрочной перспективе",
     [
         "Публичный сайт ВУЗа (4–5 нед.).",
         "Электронная библиотека и УМК (3–4 нед.).",
         "Электронный журнал посещаемости (2 нед.).",
         "Видеостудия и стриминг (3–4 нед.).",
     ]),
    ("P2 — желательные улучшения",
     [
         "Календарное расписание (2 нед.).",
         "Интеграция с ИИВ/GSP (2 нед.).",
     ]),
]


# ---------------------------------------------------------------------------
# Сборка документа
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Титул
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ\nо соответствии платформы дистанционного обучения\n"
                        "требованиям государственного чек-листа\n"
                        "«Masofaviy ta'limni joriy etilishini o'rganish (OTM)»")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    add_paragraph(doc, "")
    add_paragraph(doc, "Платформа: ndktu-student-face-platform",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Дата подготовки: 10 мая 2026 года",
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "")

    # 1. Введение
    add_heading(doc, "1. Введение", level=1)
    add_paragraph(
        doc,
        "Настоящий отчёт подготовлен в целях оценки соответствия программной "
        "платформы ndktu-student-face-platform требованиям, изложенным в чек-листе "
        "«Masofaviy ta'limni joriy etilishini o'rganish (OTM)». Документ содержит "
        "сводную таблицу соответствия по 25 показателям, описание реализованной "
        "функциональности, перечень отсутствующих компонентов, а также рекомендации "
        "по доработке и приоритизированную дорожную карту."
    )
    add_paragraph(
        doc,
        "Архитектура платформы основана на следующем стеке технологий:"
    )
    add_bullets(doc, [
        "Backend: FastAPI (Python 3.12, async), SQLAlchemy 2, PostgreSQL 17, Alembic.",
        "Frontend: React 19 + Vite + TypeScript + Tailwind CSS 4, React Query, Zod.",
        "Микросервис распознавания лиц: отдельный FastAPI-сервис на MediaPipe + face_recognition.",
        "Кеш и очереди: Redis 7.",
        "Мониторинг: Prometheus + Grafana.",
        "Развёртывание: Docker Compose, nginx-проксирование, zero-downtime deploy.",
        "Авторизация: JWT (access + refresh), RBAC через модули role / permission.",
        "Админ-панель: SQLAdmin.",
    ])

    # 2. Сводная таблица
    add_heading(doc, "2. Сводная таблица соответствия", level=1)
    add_paragraph(
        doc,
        "Цветовая маркировка статусов: зелёный — реализовано, оранжевый — реализовано "
        "частично, красный — отсутствует, серый — относится к организационным мероприятиям."
    )

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.5), Cm(0.9), Cm(5.0), Cm(2.4), Cm(5.0)]
    hdr = table.rows[0].cells
    headers = ["Раздел", "№", "Показатель", "Статус", "Комментарий"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        set_cell_shading(hdr[i], "DDDDDD")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr[i].width = widths[i]

    for section_name, num, indicator, status, comment in CHECKLIST:
        row = table.add_row().cells
        for i, value in enumerate([section_name, num, indicator, status, comment]):
            row[i].text = ""
            p = row[i].paragraphs[0]
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            row[i].width = widths[i]
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i == 3:
                color = STATUS_COLORS.get(status)
                if color:
                    run.bold = True
                    run.font.color.rgb = color
                fill_map = {
                    "Реализовано": "DFF5DF",
                    "Частично": "FCEAB6",
                    "Отсутствует": "F8D7D7",
                    "Организационный": "ECECEC",
                }
                if status in fill_map:
                    set_cell_shading(row[i], fill_map[status])

    # 3. Реализованные функции
    add_heading(doc, "3. Реализованные функции", level=1)
    add_paragraph(
        doc,
        "Ниже приведён детальный перечень функциональных компонентов, которые "
        "уже реализованы и доступны в текущей версии платформы."
    )
    for idx, (name, desc) in enumerate(IMPLEMENTED_DETAILS, start=1):
        add_heading(doc, f"3.{idx}. {name}", level=2)
        add_paragraph(doc, desc)

    # 4. Отсутствующие функции
    add_heading(doc, "4. Отсутствующие функции", level=1)
    add_paragraph(
        doc,
        "По итогам аудита кодовой базы выявлены следующие функциональные "
        "пробелы относительно требований OTM:"
    )
    add_bullets(doc, MISSING_LIST)

    # 5. Рекомендации
    add_heading(doc, "5. Рекомендации по доработке", level=1)
    add_paragraph(
        doc,
        "Для полного соответствия требованиям OTM предлагается реализовать следующие "
        "доработки. Для каждой инициативы указаны рекомендуемые технологии и "
        "ориентировочная оценка трудозатрат."
    )
    for idx, (name, desc) in enumerate(RECOMMENDATIONS, start=1):
        add_heading(doc, f"5.{idx}. {name}", level=2)
        add_paragraph(doc, desc)

    # 6. Дорожная карта
    add_heading(doc, "6. Дорожная карта", level=1)
    add_paragraph(
        doc,
        "Доработки сгруппированы по приоритетам. Общий объём работ — "
        "ориентировочно 26–34 человеко-недель плюс организационные мероприятия."
    )
    for prio_name, items in ROADMAP:
        add_heading(doc, prio_name, level=2)
        add_bullets(doc, items)

    # 7. Заключение
    add_heading(doc, "7. Заключение", level=1)
    add_paragraph(
        doc,
        "Платформа ndktu-student-face-platform закрывает значительную часть "
        "требований чек-листа OTM: реализованы LMS-ядро, интеграция с HEMIS, "
        "биометрическая идентификация, RBAC, контроль знаний, статистика и "
        "учёт контингента. Для полного соответствия требуется доработать "
        "виртуальный класс, расширенный прокторинг, средства коммуникации, "
        "публичный сайт ВУЗа и электронную библиотеку, а также пройти "
        "организационные процедуры (регистрация в reestr.uz и экспертиза DUK)."
    )
    add_paragraph(
        doc,
        "Рекомендуется приступить к работам приоритета P0 и параллельно начать "
        "оформление документов для регистрации в Едином реестре и прохождения "
        "экспертизы кибербезопасности."
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Отчёт сохранён: {OUTPUT}")


if __name__ == "__main__":
    main()
